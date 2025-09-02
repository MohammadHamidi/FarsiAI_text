#!/usr/bin/env python3
from docx import Document
import os

# Test if we can read our reference document
ref_path = '/app/outputs/reference_farsi.docx'
if os.path.exists(ref_path):
    print('Reading reference document...')
    doc = Document(ref_path)

    print(f'Document has {len(doc.paragraphs)} paragraphs')

    # Check paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f'Paragraph {i}: "{para.text}" (style: {para.style.name})')
            # Check runs in paragraph
            for j, run in enumerate(para.runs):
                if run.text.strip():
                    print(f'  Run {j}: "{run.text}" (bold: {run.bold}, italic: {run.italic})')

    # Check if our custom styles exist
    style_names = [style.name for style in doc.styles]
    print(f'\nAvailable styles: {len(style_names)}')

    # Look for our custom styles
    custom_styles = [name for name in style_names if 'Strong' in name or 'Emphasis' in name]
    print(f'Custom styles found: {custom_styles}')

else:
    print('Reference document not found')

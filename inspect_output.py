#!/usr/bin/env python3
from docx import Document
import os

output_path = '/tmp/test_output.docx'
if os.path.exists(output_path):
    print('Reading generated DOCX...')
    doc = Document(output_path)

    print(f'Document has {len(doc.paragraphs)} paragraphs')

    # Check paragraphs and their styles
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f'Paragraph {i}: "{para.text}" (style: {para.style.name})')

            # Check runs in paragraph for formatting
            for j, run in enumerate(para.runs):
                if run.text.strip():
                    print(f'  Run {j}: "{run.text}" (bold: {run.bold}, italic: {run.italic}, style: {run.style.name if run.style else None})')

    # Check if any custom styles are being used
    all_styles = []
    for para in doc.paragraphs:
        all_styles.append(para.style.name)
        for run in para.runs:
            if run.style:
                all_styles.append(run.style.name)

    unique_styles = list(set(all_styles))
    print(f'\nStyles used in document: {unique_styles}')

else:
    print('Output file not found')

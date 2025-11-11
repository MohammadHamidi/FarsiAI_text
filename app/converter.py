import subprocess
import os
import logging
from datetime import datetime, timedelta
import glob
from pathlib import Path
import zipfile
import xml.etree.ElementTree as ET
from io import BytesIO
# Ensure these are correctly imported from your config module
# and that the config module itself is correctly set up.
from .config import OUT_DIR, ALLOWED, MAX_FILE_AGE_HOURS

# Setup logging
logger = logging.getLogger(__name__)


def create_reference_docx():
    """
    Create a proper RTL reference document for better Farsi rendering
    """
    reference_path = Path(OUT_DIR) / "reference_farsi.docx"

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.shared import OxmlElement, qn
        from docx.oxml import parse_xml

        logger.info("Creating enhanced RTL reference document")

        # Create new document
        doc = Document()

        # Configure document defaults for RTL
        section = doc.sections[0]
        section.page_height = Pt(842)  # A4 height
        section.page_width = Pt(595)   # A4 width
        section.left_margin = Pt(72)   # 1 inch
        section.right_margin = Pt(72)  # 1 inch
        section.top_margin = Pt(72)    # 1 inch
        section.bottom_margin = Pt(72) # 1 inch

        # Get styles object
        styles = doc.styles

        # Configure Normal style for RTL
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'DejaVu Sans'  # Available font with Unicode support
        normal_font.size = Pt(12)
        normal_font.complex_script = True  # Enable complex script support

        # Set RTL alignment
        normal_para = normal_style.paragraph_format
        normal_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        normal_para.line_spacing_rule = WD_LINE_SPACING.SINGLE
        normal_para.space_after = Pt(6)

        # Add RTL properties to normal style
        normal_style_element = normal_style._element
        pPr = normal_style_element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            normal_style_element.insert(0, pPr)

        # Add BiDi (bi-directional) property
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)

        # Configure heading styles
        for i in range(1, 7):
            heading_style = styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = 'DejaVu Sans'
            heading_font.complex_script = True
            heading_font.bold = True

            # Set appropriate sizes
            sizes = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 11}
            heading_font.size = Pt(sizes[i])

            # Set RTL alignment for headings
            heading_para = heading_style.paragraph_format
            heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            heading_para.space_before = Pt(12)
            heading_para.space_after = Pt(6)

            # Add RTL properties to heading style
            heading_style_element = heading_style._element
            pPr = heading_style_element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                heading_style_element.insert(0, pPr)

            bidi = pPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = OxmlElement('w:bidi')
                pPr.append(bidi)

        # Modify existing Strong character style for proper bold formatting
        try:
            # Try to get existing style first
            strong_style = styles['Strong']
        except KeyError:
            # Create new style if it doesn't exist
            strong_style = styles.add_style('Strong', WD_STYLE_TYPE.CHARACTER)

        strong_font = strong_style.font
        strong_font.bold = True
        strong_font.name = 'DejaVu Sans'
        strong_font.complex_script = True

        # Add complex script bold property
        strong_style_element = strong_style._element
        rPr = strong_style_element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            strong_style_element.insert(0, rPr)

        # Clear existing formatting and add both regular and complex script bold
        # Remove existing b and bCs elements
        for b_elem in rPr.findall(qn('w:b')):
            rPr.remove(b_elem)
        for bcs_elem in rPr.findall(qn('w:bCs')):
            rPr.remove(bcs_elem)

        # Add both regular and complex script bold
        b_element = OxmlElement('w:b')
        bCs_element = OxmlElement('w:bCs')
        rPr.append(b_element)
        rPr.append(bCs_element)

        # Modify existing Emphasis character style
        try:
            # Try to get existing style first
            emphasis_style = styles['Emphasis']
        except KeyError:
            # Create new style if it doesn't exist
            emphasis_style = styles.add_style('Emphasis', WD_STYLE_TYPE.CHARACTER)

        emphasis_font = emphasis_style.font
        emphasis_font.italic = True
        emphasis_font.name = 'DejaVu Sans'
        emphasis_font.complex_script = True

        # Add complex script italic property
        emphasis_style_element = emphasis_style._element
        rPr = emphasis_style_element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            emphasis_style_element.insert(0, rPr)

        # Clear existing formatting and add both regular and complex script italic
        # Remove existing i and iCs elements
        for i_elem in rPr.findall(qn('w:i')):
            rPr.remove(i_elem)
        for ics_elem in rPr.findall(qn('w:iCs')):
            rPr.remove(ics_elem)

        # Add both regular and complex script italic
        i_element = OxmlElement('w:i')
        iCs_element = OxmlElement('w:iCs')
        rPr.append(i_element)
        rPr.append(iCs_element)

        # Also ensure the default styles are set up correctly
        # Modify the default strong and emphasis styles that Pandoc uses
        try:
            # Pandoc uses 'Strong' for **bold** and 'Emphasis' for *italic*
            # Make sure these are properly configured
            styles['Strong'].font.bold = True
            styles['Strong'].font.complex_script = True
            styles['Emphasis'].font.italic = True
            styles['Emphasis'].font.complex_script = True
        except KeyError:
            pass

        # Configure Table styles for RTL
        try:
            table_style = styles['Table Grid']
            table_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        except KeyError:
            logger.warning("Table Grid style not found")

        # Add sample content to establish proper formatting
        doc.add_heading('نمونه سند فارسی', 0)

        p1 = doc.add_paragraph('این یک نمونه متن فارسی است که برای تنظیم قالب‌بندی RTL استفاده می‌شود.')

        doc.add_heading('عنوان اول', level=1)

        # Create paragraph with mixed formatting using our custom styles
        p2 = doc.add_paragraph()
        p2.add_run('متن عادی، ').bold = False
        p2.add_run('متن پررنگ').bold = True
        p2.add_run(' و ')
        p2.add_run('متن کج').italic = True
        p2.add_run('.')

        # Add a paragraph that uses the Strong and Emphasis character styles explicitly
        p3 = doc.add_paragraph()
        run1 = p3.add_run('متن عادی ')
        run2 = p3.add_run('پررنگ')
        run2.style = styles['Strong']  # Explicitly use Strong style
        run3 = p3.add_run(' و ')
        run4 = p3.add_run('کج')
        run4.style = styles['Emphasis']  # Explicitly use Emphasis style
        run5 = p3.add_run(' با استایل‌های شخصی‌سازی شده')

        # Add a table sample
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'

        # Fill table with RTL content
        cells = [
            ['ستون ۳', 'ستون ۲', 'ستون ۱'],
            ['سلول ۳', 'سلول ۲', 'سلول ۱']
        ]

        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell.text = cells[row_idx][cell_idx]
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Save the reference document
        doc.save(str(reference_path))
        logger.info(f"RTL reference document created successfully: {reference_path}")
        return reference_path

    except ImportError:
        logger.error("python-docx not available for creating reference document")
        return None
    except Exception as e:
        logger.error(f"Failed to create reference document: {e}")
        return None

def render_markdown(md_path: str, out_path: str, fmt: str):
    """
    Convert Markdown to the specified format using Pandoc.
    Currently focused on DOCX with RTL text justification and custom main font.

    Args:
        md_path: Path to source markdown file
        out_path: Path where output should be saved
        fmt: Output format (should be "docx" based on current ALLOWED config)

    Raises:
        ValueError: If format is not supported based on ALLOWED formats.
        subprocess.CalledProcessError: If Pandoc returns a non-zero exit code.
        subprocess.TimeoutExpired: If Pandoc command times out.
        FileNotFoundError: If Pandoc executable is not found.
        Exception: For other unexpected OS or subprocess errors during conversion.
    """
    if fmt not in ALLOWED:
        logger.error(f"Unsupported format requested: '{fmt}'. Allowed formats are: {ALLOWED}")
        raise ValueError(f"Unsupported format: '{fmt}'. Only formats in {list(ALLOWED)} are supported.")

    logger.info(f"Attempting to convert '{md_path}' to '{fmt}', outputting to '{out_path}'")

    cmd = ["pandoc", md_path, "-o", out_path]

    if fmt == "docx":
        # Create enhanced reference document
        reference_path = create_reference_docx()

        # Enhanced RTL configuration
        cmd.extend([
            # Language and direction settings - CRITICAL for RTL
            "--metadata=lang:fa",           # Persian language
            "--metadata=dir:rtl",           # Right-to-left direction
            "--metadata=documentclass:article",

            # Font configuration using available fonts
            "--variable=mainfont:DejaVu Sans",      # Available font with Unicode support
            "--variable=sansfont:DejaVu Sans",
            "--variable=monofont:DejaVu Sans Mono",

            # Document formatting
            "--variable=fontsize:12pt",
            "--variable=linestretch:1.5",
            "--variable=geometry:margin=2.54cm",    # Standard Word margins

            # RTL-specific options - CRITICAL
            "--variable=rtl:true",

            # Table of contents (if needed)
            "--toc-depth=3",

            # Better list formatting
            "--variable=indent:true",

            # Image handling
            "--variable=graphics:true",
        ])

        # Use enhanced reference document if available
        if reference_path and reference_path.exists():
            cmd.append(f"--reference-doc={reference_path}")
            logger.info(f"Using enhanced RTL reference document: {reference_path}")
        else:
            logger.warning("Enhanced reference document not available, using Pandoc defaults")

        # Additional filters for better RTL handling
        cmd.extend([
            "--wrap=preserve",           # Preserve line wrapping
        ])

        logger.info(f"Enhanced DOCX configuration with comprehensive RTL support.")

    logger.debug(f"Executing Pandoc command: \"{' '.join(cmd)}\"")

    try:
        result = subprocess.run(
            cmd,
            check=True,
            timeout=120,  # Increased timeout for complex RTL documents
            capture_output=True,
            text=True,
            encoding='utf-8'
        )
        if result.stdout and result.stdout.strip():
            logger.debug(f"Pandoc STDOUT:\n{result.stdout.strip()}")
        if result.stderr and result.stderr.strip():
            logger.warning(f"Pandoc STDERR (may contain warnings/info):\n{result.stderr.strip()}")
        logger.info(f"Successfully converted '{md_path}' to '{out_path}'")

        # Post-process DOCX to ensure all paragraphs have RTL BiDi property
        if fmt == "docx" and os.path.exists(out_path):
            try:
                apply_rtl_to_docx(out_path)
                logger.info("RTL post-processing completed successfully")
            except Exception as e:
                logger.error(f"Failed to apply RTL post-processing: {e}")
                # Don't fail the entire conversion if post-processing fails
                # The file should still be usable even without this enhancement

        return out_path
    except subprocess.CalledProcessError as e:
        logger.error(f"Pandoc command failed with exit code {e.returncode} during conversion of '{md_path}' to '{fmt}'.")
        logger.error(f"Pandoc command executed: \"{' '.join(e.cmd)}\"")
        logger.error(f"Pandoc STDOUT (on error): {e.stdout.strip() if e.stdout and e.stdout.strip() else '<empty>'}")
        logger.error(f"Pandoc STDERR (on error): {e.stderr.strip() if e.stderr and e.stderr.strip() else '<empty>'}")
        if os.path.exists(out_path):
            try: os.remove(out_path)
            except OSError as re: logger.warning(f"Failed to remove '{out_path}' after error: {re}")
        raise
    except subprocess.TimeoutExpired as e:
        logger.error(f"Pandoc command timed out after {e.timeout}s for '{md_path}' to '{fmt}'. Cmd: \"{' '.join(e.cmd)}\"")
        if os.path.exists(out_path):
            try: os.remove(out_path)
            except OSError as re: logger.warning(f"Failed to remove '{out_path}' after timeout: {re}")
        raise
    except FileNotFoundError:
        logger.critical(f"Pandoc executable ('{cmd[0]}') not found. Ensure Pandoc is installed and in PATH.")
        raise
    except Exception as e:
        logger.error(f"Unexpected error during conversion of '{md_path}' to '{fmt}': {e}", exc_info=True)
        if os.path.exists(out_path):
            try: os.remove(out_path)
            except OSError as re: logger.warning(f"Failed to remove '{out_path}' after error: {re}")
        raise

def apply_rtl_to_docx(docx_path: str):
    """
    Post-process a DOCX file to ensure all paragraphs have the RTL BiDi property.

    This function opens the DOCX file (which is a ZIP archive), extracts the
    word/document.xml file, parses it, and adds the <w:bidi/> element to all
    paragraph properties (<w:pPr>) to ensure proper RTL rendering.

    Args:
        docx_path: Path to the DOCX file to process

    Raises:
        FileNotFoundError: If the DOCX file doesn't exist
        Exception: If there's an error processing the DOCX file
    """
    if not os.path.exists(docx_path):
        logger.error(f"DOCX file not found: {docx_path}")
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    logger.info(f"Applying RTL BiDi properties to all paragraphs in: {docx_path}")

    try:
        # Define XML namespaces used in DOCX files
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        }

        # Register namespaces to preserve prefixes
        for prefix, uri in namespaces.items():
            ET.register_namespace(prefix, uri)

        # Open the DOCX file as a ZIP archive
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # Read the main document XML
            document_xml = docx_zip.read('word/document.xml')

        # Parse the XML
        tree = ET.ElementTree(ET.fromstring(document_xml))
        root = tree.getroot()

        # Counter for modified paragraphs
        modified_count = 0

        # Find all paragraph elements (<w:p>)
        for paragraph in root.findall('.//w:p', namespaces):
            # Check if paragraph properties (<w:pPr>) exist
            pPr = paragraph.find('w:pPr', namespaces)

            if pPr is None:
                # Create <w:pPr> if it doesn't exist
                pPr = ET.Element('{' + namespaces['w'] + '}pPr')
                # Insert at the beginning of the paragraph
                paragraph.insert(0, pPr)

            # Check if <w:bidi/> already exists
            bidi = pPr.find('w:bidi', namespaces)

            if bidi is None:
                # Create and add <w:bidi/> element at the beginning of <w:pPr>
                bidi = ET.Element('{' + namespaces['w'] + '}bidi')
                pPr.insert(0, bidi)
                modified_count += 1

        logger.info(f"Added RTL BiDi property to {modified_count} paragraphs")

        # Convert the modified XML tree back to bytes
        modified_xml = ET.tostring(root, encoding='utf-8', xml_declaration=True)

        # Create a new DOCX file with the modified document.xml
        # We need to read all files from the original DOCX and write them to a new one
        temp_docx_path = docx_path + '.tmp'

        with zipfile.ZipFile(docx_path, 'r') as docx_zip_read:
            with zipfile.ZipFile(temp_docx_path, 'w', zipfile.ZIP_DEFLATED) as docx_zip_write:
                # Copy all files except word/document.xml
                for item in docx_zip_read.infolist():
                    if item.filename != 'word/document.xml':
                        data = docx_zip_read.read(item.filename)
                        docx_zip_write.writestr(item, data)

                # Write the modified document.xml
                docx_zip_write.writestr('word/document.xml', modified_xml)

        # Replace the original file with the modified one
        os.replace(temp_docx_path, docx_path)

        logger.info(f"Successfully applied RTL BiDi properties to: {docx_path}")

    except FileNotFoundError:
        raise
    except Exception as e:
        logger.error(f"Error applying RTL properties to DOCX: {e}", exc_info=True)
        # Clean up temp file if it exists
        temp_docx_path = docx_path + '.tmp'
        if os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except OSError:
                pass
        raise

def cleanup_old_files():
    """
    Remove temporary files older than MAX_FILE_AGE_HOURS from OUT_DIR.
    """
    if not OUT_DIR or not os.path.isdir(OUT_DIR):
        logger.warning(f"Output directory '{OUT_DIR}' is not configured or does not exist. Skipping cleanup.")
        return 0
    try:
        cutoff_time = datetime.now() - timedelta(hours=MAX_FILE_AGE_HOURS)
    except TypeError:
        logger.error(f"Invalid MAX_FILE_AGE_HOURS: '{MAX_FILE_AGE_HOURS}'. Must be a number. Skipping cleanup.")
        return 0
    count = 0
    logger.info(f"Starting cleanup of files older than {cutoff_time.strftime('%Y-%m-%d %H:%M:%S')} in '{OUT_DIR}'")
    try:
        for entry in os.scandir(OUT_DIR):
            if entry.is_file():
                try:
                    file_path = entry.path
                    mod_time_timestamp = entry.stat().st_mtime
                    mod_time = datetime.fromtimestamp(mod_time_timestamp)
                    if mod_time < cutoff_time:
                        os.remove(file_path)
                        count += 1
                        logger.debug(f"Removed old file: '{file_path}' (modified: {mod_time.strftime('%Y-%m-%d %H:%M:%S')})")
                except OSError as e:
                    logger.warning(f"Failed to process or remove old file '{entry.path}': {e}")
                except Exception as e:
                    logger.warning(f"Unexpected error processing file '{entry.path}' for cleanup: {e}")
    except OSError as e:
        logger.error(f"Error scanning output directory '{OUT_DIR}' for cleanup: {e}")
        return 0
    if count > 0:
        logger.info(f"Cleaned up {count} old file(s) from '{OUT_DIR}'.")
    else:
        logger.info(f"No old files found to clean up in '{OUT_DIR}'.")
    return count